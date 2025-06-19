# -*- coding: utf-8 -*-
import streamlit as st
import json
import datetime
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import unicodedata
from PIL import Image
import base64

# ---------- CONFIGURACIÓN GENERAL ----------
st.set_page_config(page_title="Autoevaluación Transparencia Activa", layout="wide")

def show_logo(logo_light="TRIVIA.png", logo_dark="TRIVIA_dark.png", width=180):
    try:
        img_light = Image.open(logo_light)
        buffered_light = BytesIO()
        img_light.save(buffered_light, format="PNG")
        img_str_light = base64.b64encode(buffered_light.getvalue()).decode()

        img_dark = Image.open(logo_dark)
        buffered_dark = BytesIO()
        img_dark.save(buffered_dark, format="PNG")
        img_str_dark = base64.b64encode(buffered_dark.getvalue()).decode()

        st.markdown(
            f'''
            <style>
            .logo-switch-wrapper {{
                width: {width}px;
                max-width: 45vw;
                margin-top: 0.4em;
                margin-bottom: -1.2em;
                text-align: right;
                float: right;
                background: transparent;
            }}
            .logo-light {{ display: block; }}
            .logo-dark  {{ display: none; }}
            @media (prefers-color-scheme: dark) {{
                .logo-light {{ display: none !important; }}
                .logo-dark  {{ display: block !important; }}
            }}
            </style>
            <div class="logo-switch-wrapper">
                <img src='data:image/png;base64,{img_str_light}' class='logo-light' width='{width}px'/>
                <img src='data:image/png;base64,{img_str_dark}' class='logo-dark' width='{width}px'/>
            </div>
            ''',
            unsafe_allow_html=True
        )
    except Exception as e:
        st.warning(f"No se pudo cargar el logo: {e}")


show_logo("TRIVIA.png", width=180)

st.markdown("""
    <style>
        html, body, [class*="css"]  {
            font-family: 'Montserrat', 'Aptos', Arial, sans-serif !important;
            background-color: #FFF !important;
            color: #222 !important;
        }
        .negrita { font-weight: bold; }
        .mayus { text-transform: uppercase; }
        div[role="radiogroup"] label {
            font-size: 1.15rem !important;
            font-family: 'Montserrat', 'Aptos', Arial, sans-serif !important;
            color: #222 !important;
        }
    </style>
""", unsafe_allow_html=True)

st.title("AUTOEVALUACIÓN DE TRANSPARENCIA ACTIVA")

# ---------- CARGA DE DATOS Y MAPAS ----------
@st.cache_data
def load_data():
    with open("estructura_materias_items.json", encoding="utf-8") as f:
        materias_items = json.load(f)
    with open("estructura_indicadores_especificos_REC_FINAL.json", encoding="utf-8") as f:
        indicadores_especificos = json.load(f)
    return materias_items, indicadores_especificos

materias_items, indicadores_especificos = load_data()

materias_map = defaultdict(list)
items_id_map = dict()
items_peso_map = dict()
materia_peso_map = dict()
for mi in materias_items:
    materia = mi['Materia']
    item = mi['Ítem']
    id_ = mi['ID']
    peso = mi['Peso Materia (%)']
    materias_map[materia].append(item)
    items_id_map[(materia, item)] = id_
    items_peso_map[(materia, item)] = peso
    try:
        peso_float = float(peso)
        materia_peso_map[materia] = peso_float
    except Exception:
        continue

ESCENARIOS = [
    "1. Organismo presenta sección con antecedentes.",
    "2. Organismo indica no tener antecedentes / no aplica.",
    "3. No hay sección, pero no hay evidencia de infracción.",
    "4. No hay sección y sí hay evidencia de información faltante.",
    "5. Sección/vínculo existe pero no funciona / no muestra datos.",
]

INDICADORES_GENERALES = [
    "¿La información está disponible?",
    "¿La información está actualizada?",
    "¿La información está completa?"
]

def normalizar(txt):
    if not isinstance(txt, str):
        return ""
    return "".join(c for c in unicodedata.normalize('NFD', txt) if unicodedata.category(c) != 'Mn').lower().strip()

# --------------- SESIÓN Y DATOS GENERALES --------------------
if "evaluacion" not in st.session_state:
    st.session_state.evaluacion = dict()

MESES_ESP = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

with st.form("datos_generales"):
    col1, col2, col3 = st.columns(3)
    with col1:
        organismo = st.text_input("Nombre del organismo", "")
    with col2:
        fecha = st.date_input("Fecha de la evaluación", value=datetime.date.today(), format="DD/MM/YYYY")
    with col3:
        evaluador = st.text_input("Nombre del evaluador(a)", "")
    col4, col5 = st.columns(2)
    with col4:
        mes_eval = st.selectbox("Mes evaluado", MESES_ESP)
    with col5:
        anio_eval = st.number_input("Año evaluado", min_value=2000, max_value=2100, value=datetime.date.today().year, step=1)
    submitted = st.form_submit_button("Iniciar evaluación")
if not (organismo and evaluador and mes_eval and anio_eval):
    st.warning("Por favor, ingrese nombre del organismo, evaluador/a, mes y año evaluado para comenzar.")
    st.stop()

# ---------- MATERIAS E ÍTEMS ----------
st.header("Materias e Ítems de Transparencia Activa")
materias = list(materias_map.keys())
materia_sel = st.selectbox("Seleccione materia a evaluar", materias)
st.markdown(f"<div class='negrita'>{materia_sel}</div>", unsafe_allow_html=True)
items = materias_map[materia_sel]
item_sel = st.selectbox("Seleccione ítem a evaluar", items)
st.markdown(f"<div class='negrita'>{item_sel}</div>", unsafe_allow_html=True)
id_item = items_id_map.get((materia_sel, item_sel))
key_evaluacion = f"{materia_sel} || {item_sel}"

# ---------- FORMULARIO REACTIVO, SEGURO Y PERSISTENTE POR ÍTEM ----------
st.subheader("Evaluación del Ítem Seleccionado")

if "item_states" not in st.session_state:
    st.session_state.item_states = {}

if key_evaluacion not in st.session_state.item_states:
    st.session_state.item_states[key_evaluacion] = {
        "escenario": None,
        "ig1": None,
        "ig2": None,
        "ig3": None,
        "ie": {},
        "obs": "",
    }

state = st.session_state.item_states[key_evaluacion]

# ESCENARIO
escenario = st.radio("Escenario", ESCENARIOS, key=f"escenario_{id_item}", index=ESCENARIOS.index(state["escenario"]) if state["escenario"] in ESCENARIOS else 0)
state["escenario"] = escenario

# Secuencia IG y IE
ig1_val = ig2_val = ig3_val = None
mostrar_ig2 = mostrar_ig3 = mostrar_ie = False

if escenario.startswith("1"):
    ig1_val = st.radio(
        INDICADORES_GENERALES[0],
        ["Sí", "No"],
        key=f"ig1_{id_item}",
        index=["Sí", "No"].index(state["ig1"]) if state["ig1"] in ["Sí", "No"] else 0,
    )
    if ig1_val != state["ig1"]:
        state["ig2"] = None
        state["ig3"] = None
        state["ie"] = {}
    state["ig1"] = ig1_val

    if ig1_val == "Sí":
        mostrar_ig2 = True
        ig2_val = st.radio(
            INDICADORES_GENERALES[1],
            ["Sí", "No"],
            key=f"ig2_{id_item}",
            index=["Sí", "No"].index(state["ig2"]) if state["ig2"] in ["Sí", "No"] else 0,
        )
        if ig2_val != state["ig2"]:
            state["ig3"] = None
            state["ie"] = {}
        state["ig2"] = ig2_val

        if ig2_val == "Sí":
            mostrar_ig3 = True
            ig3_opciones = ["Sí", "No", "No es posible determinarlo"]
            ig3_val = st.radio(
                INDICADORES_GENERALES[2],
                ig3_opciones,
                key=f"ig3_{id_item}",
                index=ig3_opciones.index(state["ig3"]) if state["ig3"] in ig3_opciones else 0,
            )
            if ig3_val != state["ig3"]:
                state["ie"] = {}
            state["ig3"] = ig3_val

# IE: asegúrate que todos los ítems tengan al menos una lista vacía
lista_ie = indicadores_especificos.get(key_evaluacion)
if lista_ie is None:
    key_norm = f"{normalizar(materia_sel)} || {normalizar(item_sel)}"
    for k in indicadores_especificos.keys():
        if normalizar(k) == key_norm:
            lista_ie = indicadores_especificos[k]
            break
if lista_ie is None:
    lista_ie = []
if not lista_ie:
    st.info("Este ítem aún no tiene indicadores específicos definidos en el sistema.")
mostrar_ie = mostrar_ig3 and state["ig3"] in ["Sí", "No", "No es posible determinarlo"] and len(lista_ie) > 0

if mostrar_ie:
    st.markdown("**Indicadores Específicos**")
    for ie in lista_ie:
        old_val = state["ie"].get(ie['codigo'])
        val = st.radio(
            ie["texto"],
            ["Sí", "No", "No aplica"],
            key=f"ie_{id_item}_{ie['codigo']}",
            index=(["Sí", "No", "No aplica"].index(old_val) if old_val in ["Sí", "No", "No aplica"] else 0),
        )
        state["ie"][ie['codigo']] = val

obs_val = st.text_area("Observaciones o comentarios (opcional)", value=state["obs"], key=f"obs_{id_item}")
state["obs"] = obs_val

# --------------- VALIDACIÓN Y GUARDADO ---------------------
error = None
puede_guardar = False

if escenario.startswith("2") or escenario.startswith("3"):
    puede_guardar = True
elif escenario.startswith("4") or escenario.startswith("5"):
    puede_guardar = True
elif escenario.startswith("1"):
    if not state["ig1"]:
        error = "Debe responder IG1."
    elif state["ig1"] == "No":
        puede_guardar = True
    elif state["ig1"] == "Sí" and not state["ig2"]:
        error = "Debe responder IG2."
    elif state["ig2"] == "No":
        puede_guardar = True
    elif state["ig2"] == "Sí" and not state["ig3"]:
        error = "Debe responder IG3."
    elif state["ig3"] in ["Sí", "No", "No es posible determinarlo"]:
        if mostrar_ie and (len([v for v in state["ie"].values() if v]) < len(lista_ie)):
            error = "Debe responder todos los indicadores específicos."
        else:
            puede_guardar = True
else:
    error = "Debe seleccionar un escenario."

if st.button("Guardar ítem"):
    if not puede_guardar:
        st.error(f"No se puede guardar: {error}")
    else:
        st.session_state.evaluacion[key_evaluacion] = {
            "escenario": state["escenario"],
            "ig": [state["ig1"], state["ig2"], state["ig3"]][:3],
            "ie": [{"codigo": k, "texto": next((ie['texto'] for ie in lista_ie if ie['codigo']==k), k), "respuesta": v} for k,v in state["ie"].items()] if mostrar_ie else [],
            "obs": state["obs"]
        }
        st.success("Ítem guardado correctamente.")

# -------------------- CÁLCULO Y EXPORTACIÓN ---------------------

def calcular_cumplimiento(evaluacion, materias_items, indicadores_especificos):
    cumplimiento_materia = defaultdict(list)
    hallazgos = []
    totales_materia = defaultdict(lambda: {"cumplidos": 0, "total": 0, "peso": 0.0, "excluidos": 0, "no_eval": 0})
    materia_excluida = set()
    items_eval_map = defaultdict(list)

    for mi in materias_items:
        materia = mi['Materia']
        item = mi['Ítem']
        peso = mi['Peso Materia (%)']
        key = f"{materia} || {item}"
        ev = evaluacion.get(key, {})
        esc = ev.get('escenario', "")
        ig = ev.get('ig', [])
        ie = ev.get('ie', [])

        excluido = False
        cumplimiento_item = None
        evaluado = True

        if esc.startswith("2") or esc.startswith("3"):
            excluido = True
            evaluado = False
        elif esc.startswith("4") or esc.startswith("5"):
            cumplimiento_item = 0
        elif esc.startswith("1"):
            if ig and ig[0] == "No":
                cumplimiento_item = 0
            elif len(ig) >= 2 and ig[1] == "No":
                cumplimiento_item = 0
            elif len(ig) == 3 and ig[2] in ["Sí", "No", "No es posible determinarlo"]:
                val_ig = []
                val_map = {"Sí": 1, "No": 0, "No es posible determinarlo": 1}
                val_map_ig3 = {"Sí": 1, "No": 0.25, "No es posible determinarlo": 1}
                val_ig.append(val_map.get(ig[0], 0))
                val_ig.append(val_map.get(ig[1], 0))
                val_ig.append(val_map_ig3.get(ig[2], 0))
                cumplimiento_ig = min(val_ig)
                if ie:
                    total_ies = len(ie)
                    total_ies_validas = sum(1 for e in ie if e['respuesta'] == "Sí")
                    cumplimiento_ie = total_ies_validas / total_ies if total_ies else 1
                else:
                    cumplimiento_ie = 1
                cumplimiento_item = 0.75 * cumplimiento_ig + 0.25 * cumplimiento_ie
            else:
                cumplimiento_item = 0
        else:
            evaluado = False
            excluido = True

        if excluido:
            totales_materia[materia]['excluidos'] += 1
            items_eval_map[materia].append({"item": item, "cumplimiento": None, "evaluado": False, "obs": ev.get('obs', '')})
        else:
            if cumplimiento_item is not None:
                totales_materia[materia]['cumplidos'] += cumplimiento_item
                totales_materia[materia]['total'] += 1
                totales_materia[materia]['peso'] = float(peso) if isinstance(peso, (float, int, str)) and str(peso).replace('.','',1).isdigit() else 0
                if cumplimiento_item < 1 or esc.startswith("4") or esc.startswith("5"):
                    hallazgos.append({"item": item, "obs": ev.get('obs', '')})
                cumplimiento_materia[materia].append({"item": item, "cumplimiento": round(100 * cumplimiento_item, 1)})
                items_eval_map[materia].append({"item": item, "cumplimiento": round(100 * cumplimiento_item, 1), "evaluado": True, "obs": ev.get('obs', '')})
            else:
                items_eval_map[materia].append({"item": item, "cumplimiento": None, "evaluado": False, "obs": ev.get('obs', '')})

    total_peso_usable = 0.0
    materias_incluidas = []
    pesos_ajustados = {}
    for mat, data in totales_materia.items():
        n_items = data['total']
        n_excluidos = data['excluidos']
        if n_items == 0 and n_excluidos > 0:
            continue
        if n_items > 0:
            materias_incluidas.append(mat)
            total_peso_usable += materia_peso_map.get(mat, 0)
    for mat in materias_incluidas:
        pesos_ajustados[mat] = materia_peso_map.get(mat, 0) / total_peso_usable if total_peso_usable > 0 else 0

    total_cumplimiento = 0.0
    for mat in materias_incluidas:
        data = totales_materia[mat]
        if data['total'] > 0:
            porcentaje_mat = data['cumplidos'] / data['total']
            total_cumplimiento += porcentaje_mat * pesos_ajustados[mat]
    cumplimiento_global = round(100 * total_cumplimiento, 1) if total_peso_usable else 0.0
    return cumplimiento_global, cumplimiento_materia, items_eval_map, hallazgos

# --------- Helpers para formato Word ----------

def set_black_font(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.name = 'Aptos'

def set_header_style(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '000000')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.name = 'Aptos'
            run.bold = True
            run.font.size = Pt(13)
        paragraph.alignment = 1

def set_title_style(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0,0,0)
        run.font.name = 'Aptos'
        run.bold = True
        run.font.size = Pt(13)
    paragraph.alignment = 1  # Centrado

def set_cell_center(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = 1

def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def set_table_fit_window(table):
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), "5000") # 100%
    tblPr.append(tblW)

# -------------------- EXPORTAR WORD ---------------------

def exportar_word(organismo, fecha, evaluador, mes_eval, anio_eval, cumplimiento_global, cumplimiento_materia, items_eval_map, hallazgos):
    doc = Document()
    # --- ORIENTACIÓN HORIZONTAL Y MÁRGENES ---
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # --- TÍTULO CENTRADO, PEQUEÑO Y NEGRO ---
    titulo = doc.add_paragraph()
    run = titulo.add_run('INFORME DE AUTOEVALUACIÓN DE CUMPLIMIENTO\nEN TRANSPARENCIA ACTIVA')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0,0,0)
    run.font.name = 'Aptos'
    titulo.alignment = 1 # Centrado

    # --- ESPACIO ---
    doc.add_paragraph()

    # --- DATOS GENERALES ---
    datos = doc.add_paragraph()
    datos.paragraph_format.space_after = Pt(0)
    def add_bold(label, value):
        run1 = datos.add_run(label)
        run1.bold = True
        run1.font.name = "Aptos"
        run2 = datos.add_run(str(value))
        run2.bold = False
        run2.font.name = "Aptos"
    add_bold("Organismo: ", organismo)
    datos.add_run("\n")
    add_bold("Mes evaluado: ", f"{mes_eval} {anio_eval}")
    datos.add_run("\n")
    add_bold("Fecha: ", fecha.strftime('%d/%m/%Y'))
    datos.add_run("\n")
    add_bold("Evaluador(a): ", evaluador)
    datos.add_run("\n")
    add_bold("Cumplimiento Global Observado: ", f"{cumplimiento_global:.1f} %")

    # --- ESPACIO ANTES DE TABLAS ---
    doc.add_paragraph()

    # --- TABLA POR MATERIA ---
    hmat = doc.add_paragraph("CUMPLIMIENTO POR MATERIA")
    set_title_style(hmat)
    tmat = doc.add_table(rows=1, cols=2)
    tmat.style = 'Table Grid'
    tmat.allow_autofit = True
    set_table_fit_window(tmat)
    widths = [Inches(7), Inches(2.5)]
    set_column_widths(tmat, widths)
    tmat.cell(0,0).text = "Materia"
    tmat.cell(0,1).text = "%"
    set_header_style(tmat.cell(0,0))
    set_header_style(tmat.cell(0,1))
    for mat, items in items_eval_map.items():
        if not any(i.get("evaluado") for i in items):
            row = tmat.add_row().cells
            row[0].text = mat
            row[1].text = "No se evalúa"
        else:
            avg = round(sum(i.get("cumplimiento", 0) for i in items if i.get("evaluado")) / max(len([i for i in items if i.get("evaluado")]),1), 1)
            row = tmat.add_row().cells
            row[0].text = mat
            row[1].text = f"{avg:.1f}"
        set_black_font(row[0])
        set_black_font(row[1])
        set_cell_center(row[1])
    doc.add_paragraph()

    # --- TABLA POR ÍTEM ---
    hit = doc.add_paragraph("CUMPLIMIENTO POR ÍTEM")
    set_title_style(hit)
    titem = doc.add_table(rows=1, cols=2)
    titem.style = 'Table Grid'
    titem.allow_autofit = True
    set_table_fit_window(titem)
    widths = [Inches(7), Inches(2.5)]
    set_column_widths(titem, widths)
    titem.cell(0,0).text = "Ítem"
    titem.cell(0,1).text = "%"
    set_header_style(titem.cell(0,0))
    set_header_style(titem.cell(0,1))
    for mat, items in items_eval_map.items():
        for item_data in items:
            row = titem.add_row().cells
            row[0].text = item_data['item']
            if item_data['evaluado']:
                row[1].text = f"{item_data['cumplimiento']:.1f}"
            else:
                row[1].text = "No se evalúa"
            set_black_font(row[0])
            set_black_font(row[1])
            set_cell_center(row[1])
    doc.add_paragraph()

    # --- HALLAZGOS DE INCUMPLIMIENTO ---
    hhall = doc.add_paragraph("HALLAZGOS DE INCUMPLIMIENTO POR ÍTEM")
    set_title_style(hhall)
    th = doc.add_table(rows=1, cols=2)
    th.style = 'Table Grid'
    th.allow_autofit = True
    set_table_fit_window(th)
    th.cell(0,0).text = "Ítem"
    th.cell(0,1).text = "Observaciones"
    set_header_style(th.cell(0,0))
    set_header_style(th.cell(0,1))
    for mat, items in items_eval_map.items():
        for item_data in items:
            if item_data['evaluado'] and item_data.get("cumplimiento", 100) < 100:
                row = th.add_row().cells
                row[0].text = item_data['item']
                row[1].text = item_data['obs']
            elif not item_data['evaluado']:
                row = th.add_row().cells
                row[0].text = item_data['item']
                row[1].text = "No se evalúa"
            set_black_font(row[0])
            set_black_font(row[1])
    doc.add_paragraph()

    # --- DETALLE DE INDICADORES GENERALES Y ESPECÍFICOS EN INCUMPLIMIENTO ---
    hdet = doc.add_paragraph('DETALLE DE INDICADORES GENERALES Y ESPECÍFICOS EN INCUMPLIMIENTO')
    set_title_style(hdet)
    tind = doc.add_table(rows=1, cols=3)
    tind.style = 'Table Grid'
    tind.allow_autofit = True
    set_table_fit_window(tind)
    # 45% del ancho para el código (ajusta si lo deseas)
    widths = [Inches(4), Inches(5), Inches(7)]
    set_column_widths(tind, widths)
    tind.cell(0,0).text = "Materia/Ítem"
    tind.cell(0,1).text = "Código"
    tind.cell(0,2).text = "Texto indicador"
    for c in range(3):
        set_header_style(tind.cell(0,c))
    for key, datos in st.session_state.evaluacion.items():
        materia, item = key.split(' || ')
        # IG
        if "ig" in datos and datos["ig"]:
            for idx, val in enumerate(datos["ig"]):
                if val == "No":
                    row = tind.add_row().cells
                    row[0].text = f"{materia} / {item}"
                    row[1].text = f"IG{idx+1}"
                    row[2].text = INDICADORES_GENERALES[idx]
                    for c in range(3):
                        set_black_font(row[c])
        # IE
        if "ie" in datos and datos["ie"]:
            for ie in datos["ie"]:
                if ie["respuesta"] == "No":
                    row = tind.add_row().cells
                    row[0].text = f"{materia} / {item}"
                    row[1].text = ie["codigo"]
                    row[2].text = ie["texto"]
                    for c in range(3):
                        set_black_font(row[c])
    doc.add_paragraph()

    # PIE DE PÁGINA
    section = doc.sections[-1]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "La App utilizada para esta Autoevaluación de Cumplimiento es un desarrollo de TRIVIA Capacitaciones"
    for run in paragraph.runs:
        run.font.name = 'Aptos'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0,0,0)
    paragraph.alignment = 1 # Centrado

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --------------- EXPORTAR INFORME ---------------------

st.header("Exportar informe")
if st.button("Generar y descargar informe Word"):
    cumplimiento_global, cumplimiento_materia, items_eval_map, hallazgos = calcular_cumplimiento(
        st.session_state.evaluacion, materias_items, indicadores_especificos)
    buffer = exportar_word(
        organismo,
        fecha,
        evaluador,
        mes_eval,
        anio_eval,
        cumplimiento_global,
        cumplimiento_materia,
        items_eval_map,
        hallazgos
    )
    st.download_button(
        label="Descargar informe Word",
        data=buffer,
        file_name=f"Informe_Autoevaluacion_TA_{organismo}_{fecha.strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )